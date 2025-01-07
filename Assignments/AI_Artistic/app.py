from docx import Document

# Content for the Word document
content = """
### AI Artistic Style Service Assignment

#### **1. Dockerized AI Artistic Style Service**
1. **Pulling the Docker Image:**
   - Command:
     ```bash
     docker pull urmsandeep/ai-artistic-style-service
     ```
   - **Screenshot Placeholder:**
     *Insert a screenshot of the command output here.*

2. **Running the Container:**
   - Command:
     ```bash
     docker run -d -p 5001:5001 urmsandeep/ai-artistic-style-service
     ```
   - **Screenshot Placeholder:**
     *Insert a screenshot showing the running container here.*

#### **2. API Testing**
1. **Using Curl to Test the API Endpoint:**
   - Command:
     ```bash
     curl -X POST http://127.0.0.1:5001/styleTransfer -F "image=@input.jpg" --output out.jpg
     ```
   - **Screenshot Placeholder:**
     *Insert a screenshot of the curl command execution here.*

2. **Input and Output Images:**
   - **Input Image Placeholder:**
     *Insert the input image here.*
   - **Output Image Placeholder:**
     *Insert the output image here.*

#### **3. Container Orchestration**
1. **Docker Compose File:**
   - Example Content:
     ```yaml
     version: '3.8'
     services:
       ai-artistic-style-service:
         image: urmsandeep/ai-artistic-style-service
         container_name: ai-artistic-style-service2
         ports:
           - "5001:5001"
         restart: always
     ```
   - **Screenshot Placeholder:**
     *Insert a screenshot of the Docker Compose file content here.*

2. **Running Docker Compose:**
   - Command:
     ```bash
     docker-compose up -d
     ```
   - **Screenshot Placeholder:**
     *Insert a screenshot showing all containers running here.*

#### **4. API Monitoring**
1. **Prometheus Configuration:**
   - File Content:
     ```yaml
     scrape_configs:
       - job_name: 'ai-artistic-style-service'
         static_configs:
           - targets: ['ai-artistic-style-service:5001']
     ```
   - **Screenshot Placeholder:**
     *Insert a screenshot of the Prometheus configuration file here.*

2. **Grafana Dashboard:**
   - **Screenshot Placeholder:**
     *Insert a screenshot of Grafana showing API performance metrics here.*

#### **5. CI/CD Pipeline**
1. **Jenkinsfile Content:**
   - Example Code:
     ```groovy
     pipeline {
         agent any
         stages {
             stage('Pull Docker Image') {
                 steps {
                     script {
                         sh 'docker pull urmsandeep/ai-artistic-style-service:latest'
                     }
                 }
             }
             stage('Run Tests') {
                 steps {
                     script {
                         sh 'docker run --rm -p 5001:5001 urmsandeep/ai-artistic-style-service:latest pytest tests/'
                     }
                 }
             }
             stage('Deploy Service') {
                 steps {
                     script {
                         sh 'docker-compose down && docker-compose up -d'
                     }
                 }
             }
         }
     }
     ```
   - **Screenshot Placeholder:**
     *Insert a screenshot of the Jenkinsfile content here.*

2. **Pipeline Output:**
   - **Screenshot Placeholder:**
     *Insert a screenshot of the Jenkins pipeline execution here.*

#### **6. Resilience Testing**
1. **Simulating Container Failure:**
   - Command:
     ```bash
     docker stop <container_id>
     ```
   - **Screenshot Placeholder:**
     *Insert a screenshot showing the stopped container here.*

2. **Verifying Automatic Restart:**
   - **Screenshot Placeholder:**
     *Insert a screenshot showing the container restart logs here.*

---

### Submission Checklist
- **Code Files:**
  - Docker Compose file.
  - Prometheus configuration.
  - Jenkinsfile.
- **Screenshots:**
  - Docker command outputs.
  - Input and output images.
  - Prometheus and Grafana monitoring.
  - Jenkins pipeline execution.
  - Resilience testing results.
"""

# Create a Word document
doc = Document()
doc.add_heading("AI Artistic Style Service Assignment", level=1)

# Add content to the document
for line in content.split("\n"):
    if line.startswith("### "):  # Main heading
        doc.add_heading(line.replace("### ", ""), level=2)
    elif line.startswith("#### "):  # Subheading
        doc.add_heading(line.replace("#### ", ""), level=3)
    elif line.startswith("1.") or line.startswith("2."):  # List item
        doc.add_paragraph(line, style="List Number")
    elif line.startswith("- "):  # Sub-list item
        doc.add_paragraph(line, style="List Bullet")
    elif line.strip():  # Regular text
        doc.add_paragraph(line)

# Save the document
file_path = "/mnt/data/AI_Artistic_Style_Assignment.docx"
doc.save(file_path)
file_path
